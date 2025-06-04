import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def is_rtl(text):
    # Check for Arabic/Urdu Unicode range
    return any('\u0600' <= c <= '\u06FF' for c in text)

def set_rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def convert_pdf_to_docx(pdf_path, docx_path):
    try:
        doc = fitz.open(pdf_path)
        document = Document()
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for b in blocks:
                if b['type'] == 0:  # text
                    for line in b["lines"]:
                        para = document.add_paragraph()
                        for span in line["spans"]:
                            run = para.add_run(span["text"])
                            # Font name
                            try:
                                run.font.name = span["font"]
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), span["font"])
                            except Exception:
                                pass
                            # Font size
                            try:
                                run.font.size = Pt(span["size"])
                            except Exception:
                                pass
                            # Font color
                            try:
                                if "color" in span:
                                    color = span["color"]
                                    if isinstance(color, int):
                                        r = (color >> 16) & 0xFF
                                        g = (color >> 8) & 0xFF
                                        b = color & 0xFF
                                        run.font.color.rgb = RGBColor(r, g, b)
                            except Exception:
                                pass
                            # Bold/italic
                            run.bold = span.get("flags", 0) & 2 != 0
                            run.italic = span.get("flags", 0) & 1 != 0
                        # RTL
                        if is_rtl(''.join([s["text"] for s in line["spans"]])):
                            set_rtl(para)
                elif b['type'] == 1:  # image
                    try:
                        img = b['image']
                        xref = b['image']
                        pix = fitz.Pixmap(doc, xref)
                        img_path = 'temp_img.png'
                        pix.save(img_path)
                        document.add_picture(img_path)
                        os.remove(img_path)
                    except Exception:
                        pass
                # Table and other elements can be added here
        document.save(docx_path)
        return True, None
    except Exception as e:
        return False, str(e) 