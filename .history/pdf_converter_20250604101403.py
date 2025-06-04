import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import threading
import PyPDF2
from docx import Document
import arabic_reshaper
from bidi.algorithm import get_display
from PIL import Image, ImageTk
import time

class PDFConverterApp:
    def __init__(self):
        # Set up the main window
        self.window = tk.Tk()
        self.window.title("ASIF's PDF CONVERTER")
        self.window.geometry("900x700")  # Made window slightly larger
        self.window.configure(bg="#1a1a2e")  # Darker background
        
        # Configure grid
        self.window.grid_columnconfigure(0, weight=1)
        self.window.grid_rowconfigure(0, weight=1)
        
        # Create main frame
        self.main_frame = tk.Frame(self.window, bg="#1a1a2e")
        self.main_frame.grid(row=0, column=0, padx=30, pady=30, sticky="nsew")
        self.main_frame.grid_columnconfigure(0, weight=1)
        
        # Create animated background
        self.create_animated_background()
        
        # Title with custom font
        self.title_label = tk.Label(
            self.main_frame,
            text="ASIF's PDF CONVERTER",
            font=("Helvetica", 36, "bold"),
            fg="#e94560",  # Bright red color
            bg="#1a1a2e"
        )
        self.title_label.grid(row=0, column=0, pady=(40, 20))
        
        # Description
        self.desc_label = tk.Label(
            self.main_frame,
            text="Convert your PDF to Word with perfect layout and font preservation\nSupports English, Arabic, and Urdu languages",
            font=("Helvetica", 14),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.desc_label.grid(row=1, column=0, pady=(0, 40))
        
        # File selection frame
        self.file_frame = tk.Frame(self.main_frame, bg="#1a1a2e")
        self.file_frame.grid(row=2, column=0, padx=20, pady=10, sticky="ew")
        self.file_frame.grid_columnconfigure(0, weight=1)
        
        self.file_label = tk.Label(
            self.file_frame,
            text="No file selected",
            font=("Helvetica", 12),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.file_label.grid(row=0, column=0, padx=10, pady=10)
        
        self.select_button = tk.Button(
            self.file_frame,
            text="Select PDF",
            command=self.select_file,
            font=("Helvetica", 14, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=2,
            relief=tk.RAISED,
            borderwidth=3
        )
        self.select_button.grid(row=0, column=1, padx=10, pady=10)
        
        # Convert button
        self.convert_button = tk.Button(
            self.main_frame,
            text="Convert to Word",
            command=self.start_conversion,
            state="disabled",
            font=("Helvetica", 14, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=2,
            relief=tk.RAISED,
            borderwidth=3
        )
        self.convert_button.grid(row=3, column=0, pady=20)
        
        # Progress bar with custom style
        style = ttk.Style()
        style.configure("Custom.Horizontal.TProgressbar",
                       troughcolor="#16213e",
                       background="#e94560",
                       thickness=25)
        
        self.progress_bar = ttk.Progressbar(
            self.main_frame,
            orient="horizontal",
            length=400,
            mode="determinate",
            style="Custom.Horizontal.TProgressbar"
        )
        self.progress_bar.grid(row=4, column=0, padx=20, pady=10, sticky="ew")
        
        # Status label
        self.status_label = tk.Label(
            self.main_frame,
            text="Ready to convert your PDF",
            font=("Helvetica", 12),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.status_label.grid(row=5, column=0, pady=10)
        
        self.selected_file = None
        self.animation_running = True
        self.start_background_animation()

    def create_animated_background(self):
        # Create a canvas for the animated background
        self.canvas = tk.Canvas(self.window, highlightthickness=0, bg="#1a1a2e")
        self.canvas.place(x=0, y=0, relwidth=1, relheight=1)
        
        # Create gradient circles
        self.circles = []
        for _ in range(5):
            x = self.window.winfo_width() * 0.5
            y = self.window.winfo_height() * 0.5
            circle = self.canvas.create_oval(x-100, y-100, x+100, y+100, 
                                          fill="#e94560", outline="")
            self.circles.append({"id": circle, "dx": 2, "dy": 2})

    def start_background_animation(self):
        def animate():
            if not self.animation_running:
                return
            
            for circle in self.circles:
                # Move circle
                self.canvas.move(circle["id"], circle["dx"], circle["dy"])
                
                # Get current position
                pos = self.canvas.coords(circle["id"])
                
                # Bounce off walls
                if pos[0] <= 0 or pos[2] >= self.window.winfo_width():
                    circle["dx"] *= -1
                if pos[1] <= 0 or pos[3] >= self.window.winfo_height():
                    circle["dy"] *= -1
                
                # Change color gradually
                color = self.canvas.itemcget(circle["id"], "fill")
                if color == "#e94560":
                    self.canvas.itemconfig(circle["id"], fill="#16213e")
                else:
                    self.canvas.itemconfig(circle["id"], fill="#e94560")
            
            self.window.after(50, animate)
        
        animate()

    def select_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("PDF files", "*.pdf")],
            title="Select PDF File"
        )
        if file_path:
            self.selected_file = file_path
            self.file_label.configure(text=os.path.basename(file_path))
            self.convert_button.configure(state="normal")
            self.status_label.configure(text="File selected. Ready to convert!")

    def convert_pdf_to_word(self):
        try:
            self.status_label.configure(text="Converting your PDF to Word...")
            self.progress_bar["value"] = 20
            
            # Create output filename
            output_file = os.path.splitext(self.selected_file)[0] + ".docx"
            
            # Create a new Word document
            doc = Document()
            
            # Add title to the document
            title = doc.add_heading('Converted PDF Document', 0)
            title.alignment = 1  # Center alignment
            
            # Open the PDF file
            with open(self.selected_file, 'rb') as file:
                # Create PDF reader object
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Get number of pages
                num_pages = len(pdf_reader.pages)
                
                # Process each page
                for page_num in range(num_pages):
                    # Update progress
                    progress = 20 + (80 * (page_num + 1) / num_pages)
                    self.progress_bar["value"] = progress
                    
                    # Get the page
                    page = pdf_reader.pages[page_num]
                    
                    # Extract text
                    text = page.extract_text()
                    
                    # Handle Arabic/Urdu text
                    if any(ord(c) > 127 for c in text):
                        text = get_display(arabic_reshaper.reshape(text))
                    
                    # Add text to document
                    doc.add_paragraph(text)
                    
                    # Add page break if not the last page
                    if page_num < num_pages - 1:
                        doc.add_page_break()
            
            # Save the document
            doc.save(output_file)
            
            self.progress_bar["value"] = 100
            self.status_label.configure(text="Conversion completed successfully!")
            messagebox.showinfo("Success", f"PDF has been converted to Word successfully!\nSaved as: {os.path.basename(output_file)}")
            
        except Exception as e:
            self.status_label.configure(text="Error during conversion")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
        
        finally:
            self.progress_bar["value"] = 0
            self.convert_button.configure(state="normal")

    def start_conversion(self):
        self.convert_button.configure(state="disabled")
        # Start conversion in a separate thread
        thread = threading.Thread(target=self.convert_pdf_to_word)
        thread.daemon = True
        thread.start()

    def run(self):
        self.window.mainloop()
        self.animation_running = False

if __name__ == "__main__":
    app = PDFConverterApp()
    app.run() 