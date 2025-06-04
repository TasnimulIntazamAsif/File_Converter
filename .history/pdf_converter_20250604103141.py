import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import threading
import PyPDF2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
import arabic_reshaper
from bidi.algorithm import get_display
from PIL import Image, ImageTk, ImageDraw, ImageFont # Import ImageFont for custom font drawing if needed for icons
import time
from datetime import datetime

# --- Icon Assets (Simple Placeholder Icons) ---
# In a real application, you'd load proper icon files (e.g., .png, .ico)
def create_icon(size, text, bg_color, fg_color):
    try:
        # Try to load a generic sans-serif font, or fall back to default
        font = ImageFont.truetype("arial.ttf", int(size * 0.6))
    except IOError:
        font = ImageFont.load_default()

    img = Image.new('RGBA', (size, size), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    draw.ellipse([0, 0, size, size], fill=bg_color)
    
    # Use textbbox instead of textsize
    bbox = draw.textbbox((0, 0), text, font=font)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    
    x = (size - text_width) / 2
    y = (size - text_height) / 2
    draw.text((x, y), text, fill=fg_color, font=font)
    return ImageTk.PhotoImage(img)

class PDFConverterApp:
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("ASIF's PDF CONVERTER")
        self.window.geometry("900x700")
        self.window.configure(bg="#1a1a2e")

        # Create and set application icon
        self.create_app_icon()

        # Configure grid for main window
        self.window.grid_columnconfigure(0, weight=1)
        self.window.grid_rowconfigure(0, weight=1)

        # Create main frame
        self.main_frame = tk.Frame(self.window, bg="#1a1a2e")
        self.main_frame.grid(row=0, column=0, padx=30, pady=30, sticky="nsew")
        self.main_frame.grid_columnconfigure(0, weight=1)

        # Create animated background (keep it as is for visual flair)
        self.create_animated_background()

        # Title with custom font and icon
        self.title_frame = tk.Frame(self.main_frame, bg="#1a1a2e")
        self.title_frame.grid(row=0, column=0, pady=(40, 20))

        # Create PDF icon for title
        self.create_title_icon()

        self.title_label = tk.Label(
            self.title_frame,
            text="ASIF's PDF CONVERTER",
            font=("Helvetica", 36, "bold"),
            fg="#e94560",
            bg="#1a1a2e"
        )
        self.title_label.pack(side="right", padx=10)

        # --- Navigation Bar (Home, About) ---
        self.nav_frame = tk.Frame(self.main_frame, bg="#16213e", height=50)
        self.nav_frame.grid(row=1, column=0, sticky="ew", pady=(0, 20))
        self.nav_frame.grid_columnconfigure((0, 1, 2), weight=1) # Center buttons

        self.home_icon = create_icon(30, "🏠", "#e94560", "white")
        self.about_icon = create_icon(30, "ℹ️", "#e94560", "white")

        self.home_button = tk.Button(
            self.nav_frame,
            image=self.home_icon,
            text=" Home", # Add text for better usability
            compound="left",
            command=lambda: self.tab_control.select(self.convert_tab), # Switch to convert tab
            font=("Helvetica", 12, "bold"),
            bg="#16213e", fg="white", activebackground="#e94560", activeforeground="white",
            relief="flat", bd=0, padx=15, pady=5
        )
        self.home_button.grid(row=0, column=0, padx=10)

        self.about_button = tk.Button(
            self.nav_frame,
            image=self.about_icon,
            text=" About",
            compound="left",
            command=self.show_about_info,
            font=("Helvetica", 12, "bold"),
            bg="#16213e", fg="white", activebackground="#e94560", activeforeground="white",
            relief="flat", bd=0, padx=15, pady=5
        )
        self.about_button.grid(row=0, column=2, padx=10) # Place at the end

        # Create tab control
        self.tab_control = ttk.Notebook(self.main_frame)
        self.tab_control.grid(row=2, column=0, sticky="ew", pady=(0, 20)) # Updated row

        # Create Convert PDF tab with icon
        self.convert_tab = tk.Frame(self.tab_control, bg="#1a1a2e")
        self.tab_control.add(self.convert_tab, text="📄 INSERT PDF FILE")

        # Create Recent Files tab with icon
        self.recent_tab = tk.Frame(self.tab_control, bg="#1a1a2e")
        self.tab_control.add(self.recent_tab, text="📋 Recent Files")

        # Configure tab style
        style = ttk.Style()
        style.configure("TNotebook", background="#1a1a2e", borderwidth=0)
        style.configure("TNotebook.Tab",
                        background="#16213e",
                        foreground="white",
                        padding=[10, 5],
                        font=("Helvetica", 10, "bold"))
        style.map("TNotebook.Tab",
                  background=[("selected", "#e94560")],
                  foreground=[("selected", "white")])

        # Setup Convert PDF tab content
        self.setup_convert_tab()

        # Setup Recent Files tab content
        self.setup_recent_tab()

        # Initialize recent files list
        self.recent_files = []
        self.max_recent_files = 10

        self.selected_file = None
        self.animation_running = True
        self.start_background_animation()

    def create_app_icon(self):
        icon_size = 32
        icon = Image.new('RGBA', (icon_size, icon_size), (0, 0, 0, 0))
        draw = ImageDraw.Draw(icon)
        draw.rectangle([2, 2, icon_size-2, icon_size-2], fill="#e94560")
        try:
            font = ImageFont.truetype("arial.ttf", 12)
            bbox = draw.textbbox((8, 8), "PDF", font=font)
            draw.text((8, 8), "PDF", fill="white", font=font)
        except IOError:
            # Fallback if font not found
            bbox = draw.textbbox((8, 8), "PDF")
            draw.text((8, 8), "PDF", fill="white")
        self.app_icon = ImageTk.PhotoImage(icon)
        self.window.iconphoto(True, self.app_icon)

    def create_title_icon(self):
        icon_size = 48
        icon = Image.new('RGBA', (icon_size, icon_size), (0, 0, 0, 0))
        draw = ImageDraw.Draw(icon)
        draw.rectangle([2, 2, icon_size-2, icon_size-2], fill="#e94560")
        try:
            font = ImageFont.truetype("arial.ttf", 20)
            bbox = draw.textbbox((8, 8), "PDF", font=font)
            draw.text((8, 8), "PDF", fill="white", font=font)
        except IOError:
            # Fallback if font not found
            bbox = draw.textbbox((8, 8), "PDF")
            draw.text((8, 8), "PDF", fill="white")
        self.title_icon = ImageTk.PhotoImage(icon)
        icon_label = tk.Label(
            self.title_frame,
            image=self.title_icon,
            bg="#1a1a2e"
        )
        icon_label.pack(side="left", padx=10)

    def setup_convert_tab(self):
        self.desc_label = tk.Label(
            self.convert_tab,
            text="📄 Convert your PDF to Word with a focus on text extraction.\n"
                 "Supports English, Arabic, and Urdu languages for text content.",
            font=("Helvetica", 14),
            fg="#ffffff",
            bg="#1a1a2e",
            wraplength=700 # Add wraplength for better text display
        )
        self.desc_label.pack(pady=(20, 40))

        # File selection frame
        self.file_frame = tk.Frame(self.convert_tab, bg="#1a1a2e")
        self.file_frame.pack(fill="x", padx=20, pady=10)

        self.file_label = tk.Label(
            self.file_frame,
            text="📄 No file selected",
            font=("Helvetica", 12),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.file_label.pack(side="left", padx=10, pady=10)

        self.select_button = tk.Button(
            self.file_frame,
            text="📂 Select PDF",
            command=self.select_file,
            font=("Helvetica", 14, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=2,
            relief=tk.RAISED,
            borderwidth=3
        )
        self.select_button.pack(side="right", padx=10, pady=10)

        # Convert button with icon
        self.convert_button = tk.Button(
            self.convert_tab,
            text="🔄 Convert to Word",
            command=self.start_conversion,
            state="disabled",
            font=("Helvetica", 14, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=20, # Increased width
            height=2,
            relief=tk.RAISED,
            borderwidth=3
        )
        self.convert_button.pack(pady=20)

        # Progress bar with custom style
        style = ttk.Style()
        style.configure("Custom.Horizontal.TProgressbar",
                        troughcolor="#16213e",
                        background="#e94560",
                        thickness=25)

        self.progress_bar = ttk.Progressbar(
            self.convert_tab,
            orient="horizontal",
            length=400,
            mode="determinate",
            style="Custom.Horizontal.TProgressbar"
        )
        self.progress_bar.pack(pady=10)

        # Status label with icon
        self.status_label = tk.Label(
            self.convert_tab,
            text="📄 Ready to convert your PDF",
            font=("Helvetica", 12),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.status_label.pack(pady=10)

    def setup_recent_tab(self):
        self.recent_frame = tk.Frame(self.recent_tab, bg="#1a1a2e")
        self.recent_frame.pack(fill="both", expand=True, padx=20, pady=20)

        self.recent_label = tk.Label(
            self.recent_frame,
            text="📋 Recently Converted Files",
            font=("Helvetica", 16, "bold"),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.recent_label.pack(pady=(0, 20))

        self.recent_listbox = tk.Listbox(
            self.recent_frame,
            bg="#16213e",
            fg="#ffffff",
            font=("Helvetica", 12),
            selectbackground="#e94560",
            selectforeground="#ffffff",
            height=10,
            borderwidth=0, # Remove border for cleaner look
            highlightthickness=0 # Remove highlight border on focus
        )
        self.recent_listbox.pack(fill="both", expand=True)

        scrollbar = tk.Scrollbar(self.recent_listbox)
        scrollbar.pack(side="right", fill="y")
        self.recent_listbox.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.recent_listbox.yview)

        button_frame = tk.Frame(self.recent_frame, bg="#1a1a2e")
        button_frame.pack(fill="x", pady=20)

        self.open_button = tk.Button(
            button_frame,
            text="📂 Open File",
            command=self.open_selected_file,
            font=("Helvetica", 12, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=1,
            relief=tk.RAISED,
            borderwidth=2
        )
        self.open_button.pack(side="left", padx=5)

        self.clear_button = tk.Button(
            button_frame,
            text="🗑️ Clear List",
            command=self.clear_recent_files,
            font=("Helvetica", 12, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=1,
            relief=tk.RAISED,
            borderwidth=2
        )
        self.clear_button.pack(side="right", padx=5)

    def add_to_recent_files(self, file_path):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        file_info = f"📄 {timestamp} - {os.path.basename(file_path)}"

        # Store the full path in a hidden list to retrieve it later
        # We'll store a tuple: (display_text, full_path)
        # This helps in opening the correct file later
        full_info = (file_info, file_path)

        # Check if the file is already in the list to avoid duplicates
        # We compare by full path
        if not any(item[1] == file_path for item in self.recent_files):
            self.recent_files.insert(0, full_info)
            if len(self.recent_files) > self.max_recent_files:
                self.recent_files.pop()

            self.recent_listbox.delete(0, tk.END)
            for display_text, _ in self.recent_files:
                self.recent_listbox.insert(tk.END, display_text)

    def open_selected_file(self):
        selection = self.recent_listbox.curselection()
        if selection:
            # Retrieve the full file path from our internal list
            selected_index = selection[0]
            _, file_path = self.recent_files[selected_index]

            if os.path.exists(file_path):
                os.startfile(file_path)
            else:
                messagebox.showerror("Error", "❌ File not found! It might have been moved or deleted.")
        else:
            messagebox.showwarning("No Selection", "Please select a file from the list to open.")

    def clear_recent_files(self):
        if messagebox.askyesno("Clear Recent Files", "Are you sure you want to clear the recent files list?"):
            self.recent_files.clear()
            self.recent_listbox.delete(0, tk.END)
            messagebox.showinfo("Cleared", "Recent files list has been cleared.")

    def select_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("PDF files", "*.pdf")],
            title="Select PDF File"
        )
        if file_path:
            self.selected_file = file_path
            self.file_label.configure(text=f"📄 {os.path.basename(file_path)}")
            self.convert_button.configure(state="normal")
            self.status_label.configure(text="📄 File selected. Ready to convert!")
        else:
            self.selected_file = None
            self.file_label.configure(text="📄 No file selected")
            self.convert_button.configure(state="disabled")
            self.status_label.configure(text="📄 Ready to convert your PDF")

    def convert_pdf_to_word(self):
        if not self.selected_file:
            messagebox.showwarning("No File", "Please select a PDF file first.")
            self.status_label.configure(text="❌ No file selected.")
            self.convert_button.configure(state="normal")
            return

        try:
            self.status_label.configure(text="🔄 Converting your PDF to Word...")
            self.progress_bar["value"] = 10 # Initial progress

            output_file = os.path.splitext(self.selected_file)[0] + "_converted.docx" # Added "_converted" to avoid overwriting
            doc = Document()

            # Add a title to the document (optional, but good for converted docs)
            doc.add_heading('Converted Document', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            with open(self.selected_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                # Set default font for the document (e.g., Calibri for English, Arabic/Urdu fonts if available)
                # This is a general setting and won't preserve original PDF fonts
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Calibri' # Default for English, common
                font.size = Pt(12)

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()

                    # Handle Arabic/Urdu text and add it as a paragraph
                    # This attempts to get the correct visual display but font embedding/layout is hard
                    if any(ord(c) > 127 for c in text): # Check for non-ASCII characters
                        # Reshape for Arabic/Urdu presentation
                        reshaped_text = arabic_reshaper.reshape(text)
                        bidi_text = get_display(reshaped_text)

                        paragraph = doc.add_paragraph(bidi_text)
                        # Set paragraph direction to RTL for Arabic/Urdu
                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Align right
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                        # Attempt to set a suitable font for Arabic/Urdu (e.g., Arial, Segoe UI)
                        # You might need to install these fonts on your system for them to work
                        run = paragraph.add_run()
                        run.text = bidi_text # Re-add text to run for font application
                        run.font.name = 'Arial' # Common font that supports Arabic/Urdu characters
                        run.font.size = Pt(12)
                        r = run._element
                        r.set(qn('w:rtl'), '1') # Set RTL property for the run

                    else:
                        doc.add_paragraph(text) # For English text, default alignment and font

                    # Add page break if not the last page to maintain page-like separation
                    if page_num < num_pages - 1:
                        doc.add_page_break()

                    progress = 10 + (90 * (page_num + 1) / num_pages)
                    self.progress_bar["value"] = progress
                    self.window.update_idletasks() # Update UI during conversion

            doc.save(output_file)

            self.add_to_recent_files(output_file)

            self.progress_bar["value"] = 100
            self.status_label.configure(text="✅ Conversion completed successfully!")
            messagebox.showinfo("Success", f"PDF has been converted to Word successfully!\nSaved as: {os.path.basename(output_file)}")

        except PyPDF2.errors.PdfReadError:
            self.status_label.configure(text="❌ Error: Invalid PDF file or corrupted.")
            messagebox.showerror("Error", "The selected file is not a valid PDF or is corrupted.")
        except Exception as e:
            self.status_label.configure(text="❌ Error during conversion")
            messagebox.showerror("Error", f"An unexpected error occurred: {str(e)}")

        finally:
            self.progress_bar["value"] = 0
            self.convert_button.configure(state="normal")
            self.window.update_idletasks() # Ensure UI is updated at the end

    def start_conversion(self):
        self.convert_button.configure(state="disabled")
        # Start conversion in a separate thread
        thread = threading.Thread(target=self.convert_pdf_to_word)
        thread.daemon = True # Allow the thread to exit with the main program
        thread.start()

    def show_about_info(self):
        about_text = (
            "ASIF's PDF Converter v1.0\n\n"
            "This application converts PDF files to Word (.docx) format.\n"
            "It aims to extract text and maintain basic page separation.\n"
            "Supports English, Arabic, and Urdu text extraction.\n\n"
            "Please note: Due to the inherent differences between PDF and Word formats,\n"
            "perfect layout, font, and image preservation can be challenging.\n"
            "Complex layouts and embedded graphics may not be accurately reproduced.\n\n"
            "Developed by ASIF"
        )
        messagebox.showinfo("About ASIF's PDF Converter", about_text)

    def create_animated_background(self):
        self.canvas = tk.Canvas(self.window, highlightthickness=0, bg="#1a1a2e")
        # Place canvas behind other widgets
        self.canvas.place(x=0, y=0, relwidth=1, relheight=1)
        
        # Create circles after window is fully initialized
        self.window.update_idletasks()
        
        self.circles = []
        for _ in range(5):
            # Initial position (center of window)
            x = self.window.winfo_width() * 0.5
            y = self.window.winfo_height() * 0.5
            circle = self.canvas.create_oval(x-100, y-100, x+100, y+100,
                                             fill="#e94560", outline="", stipple="gray50")
            self.circles.append({"id": circle, "dx": 2, "dy": 2})
        
        # Ensure canvas is behind other widgets
        self.canvas.lower()

    def start_background_animation(self):
        def animate():
            if not self.animation_running:
                return

            # Get current window dimensions
            canvas_width = self.canvas.winfo_width()
            canvas_height = self.canvas.winfo_height()

            for circle in self.circles:
                # Move circle
                self.canvas.move(circle["id"], circle["dx"], circle["dy"])
                pos = self.canvas.coords(circle["id"])

                # Bounce off walls
                if pos[0] <= 0 or pos[2] >= canvas_width:
                    circle["dx"] *= -1
                if pos[1] <= 0 or pos[3] >= canvas_height:
                    circle["dy"] *= -1

                # Change color gradually
                current_fill = self.canvas.itemcget(circle["id"], "fill")
                if current_fill == "#e94560":
                    self.canvas.itemconfig(circle["id"], fill="#16213e")
                else:
                    self.canvas.itemconfig(circle["id"], fill="#e94560")

            # Schedule next animation frame
            self.window.after(50, animate)

        # Start animation
        animate()

    def run(self):
        self.window.mainloop()
        self.animation_running = False # Stop animation thread when main loop exits

if __name__ == "__main__":
    app = PDFConverterApp()
    app.run()