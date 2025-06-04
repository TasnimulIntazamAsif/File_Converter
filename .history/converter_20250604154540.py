import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Map common PDF font names to system fonts (extend as needed)
FONT_MAP = {
    'Times-Roman': 'Times New Roman',
    'Helvetica': 'Arial',
    'Courier': 'Courier New',
    # Add more mappings as needed
}

def is_rtl(text):
    # Check for Arabic/Urdu Unicode range
    return any('\u0600' <= c <= '\u06FF' for c in text)

def set_rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_page_number_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def get_mapped_font(font_name):
    if font_name in FONT_MAP:
        return FONT_MAP[font_name]
    return font_name

def convert_pdf_to_docx(pdf_path, docx_path):
    try:
        doc = fitz.open(pdf_path)
        document = Document()
        # Set page size and margins to match the PDF's first page
        pdf_page = doc[0]
        width_pt, height_pt = pdf_page.rect.width, pdf_page.rect.height
        # 1 point = 1/72 inch
        width_inch = width_pt / 72
        height_inch = height_pt / 72
        section = document.sections[0]
        section.page_width = Inches(width_inch)
        section.page_height = Inches(height_inch)
        # Set margins to 0.5 inch as a default (can be improved by extracting from PDF if needed)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        # Add page number to footer
        set_page_number_footer(section)
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            # Sort lines by Y position for table row order
            lines = []
            for b in blocks:
                if b['type'] == 0:  # text
                    for line in b["lines"]:
                        y = line["bbox"][1]
                        lines.append((y, line))
            lines.sort(key=lambda x: x[0])
            table = document.add_table(rows=len(lines), cols=1)
            table.autofit = False
            for i, (y, line) in enumerate(lines):
                cell = table.cell(i, 0)
                para = cell.paragraphs[0]
                for span in line["spans"]:
                    run = para.add_run(span["text"])
                    # Font name mapping
                    font_name = get_mapped_font(span["font"])
                    try:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    except Exception:
                        pass
                    # Font size
                    try:
                        run.font.size = Pt(span["size"])
                    except Exception:
                        pass
                    # Font color
                    try:
                        if "color" in span:
                            color = span["color"]
                            if isinstance(color, int):
                                r = (color >> 16) & 0xFF
                                g = (color >> 8) & 0xFF
                                b = color & 0xFF
                                run.font.color.rgb = RGBColor(r, g, b)
                    except Exception:
                        pass
                    # Bold/italic
                    run.bold = span.get("flags", 0) & 2 != 0
                    run.italic = span.get("flags", 0) & 1 != 0
                # RTL
                if is_rtl(''.join([s["text"] for s in line["spans"]])):
                    set_rtl(para)
            # Add images (not positioned, but included)
            for b in blocks:
                if b['type'] == 1:  # image
                    try:
                        img = b['image']
                        xref = b['image']
                        pix = fitz.Pixmap(doc, xref)
                        img_path = 'temp_img.png'
                        pix.save(img_path)
                        document.add_picture(img_path)
                        os.remove(img_path)
                    except Exception:
                        pass
        document.save(docx_path)
        return True, None
    except Exception as e:
        return False, str(e) 