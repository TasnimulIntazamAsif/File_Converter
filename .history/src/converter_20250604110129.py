import fitz  # PyMuPDF
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image
import io
import os
import arabic_reshaper
from bidi.algorithm import get_display
import re

class PDFConverter:
    def __init__(self):
        self.supported_languages = {
            'en': r'[\x00-\x7F]',  # ASCII
            'ar': r'[\u0600-\u06FF]',  # Arabic
            'ur': r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]'  # Urdu
        }

    def detect_language(self, text):
        if not text:
            return 'en'
        
        # Check for Arabic and Urdu characters
        if re.search(self.supported_languages['ur'], text):
            return 'ur'
        elif re.search(self.supported_languages['ar'], text):
            return 'ar'
        return 'en'

    def convert(self, pdf_path, output_dir):
        # Open PDF
        pdf_document = fitz.open(pdf_path)
        doc = docx.Document()
        
        # Process each page
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Extract text blocks
            blocks = page.get_text("dict")["blocks"]
            
            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text = span["text"]
                            if not text.strip():
                                continue
                                
                            # Detect language and handle RTL text
                            lang = self.detect_language(text)
                            if lang in ['ar', 'ur']:
                                text = get_display(arabic_reshaper.reshape(text))
                            
                            # Create paragraph with proper alignment
                            p = doc.add_paragraph()
                            run = p.add_run(text)
                            
                            # Set font properties
                            font = run.font
                            font.name = span["font"]
                            font.size = Pt(span["size"])
                            
                            # Handle RTL text
                            if lang in ['ar', 'ur']:
                                p._element.set(qn('w:bidi'), '1')
                                p._element.set(qn('w:rtl'), '1')
                            
                            # Set text alignment
                            if span["flags"] & 2**4:  # Bold
                                font.bold = True
                            if span["flags"] & 2**1:  # Italic
                                font.italic = True
                            
                            # Set paragraph alignment
                            if span["flags"] & 2**3:  # Center
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif span["flags"] & 2**2:  # Right
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Extract and add images
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Convert image bytes to PIL Image
                image = Image.open(io.BytesIO(image_bytes))
                
                # Save image temporarily
                temp_img_path = f"temp_img_{img_index}.png"
                image.save(temp_img_path)
                
                # Add image to document
                doc.add_picture(temp_img_path, width=Inches(6))
                
                # Clean up temporary file
                os.remove(temp_img_path)
            
            # Add page break if not the last page
            if page_num < len(pdf_document) - 1:
                doc.add_page_break()
        
        # Save the document
        output_filename = os.path.splitext(os.path.basename(pdf_path))[0] + '.docx'
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        
        # Close PDF
        pdf_document.close() 