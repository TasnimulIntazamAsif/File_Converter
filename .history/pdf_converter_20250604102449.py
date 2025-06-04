import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import threading
import PyPDF2
from docx import Document
import arabic_reshaper
from bidi.algorithm import get_display
from PIL import Image, ImageTk, ImageDraw
import time
from datetime import datetime

class PDFConverterApp:
    def __init__(self):
        # Set up the main window
        self.window = tk.Tk()
        self.window.title("ASIF's PDF CONVERTER")
        self.window.geometry("900x700")
        self.window.configure(bg="#1a1a2e")
        
        # Create and set application icon
        self.create_app_icon()
        
        # Configure grid
        self.window.grid_columnconfigure(0, weight=1)
        self.window.grid_rowconfigure(0, weight=1)
        
        # Create main frame
        self.main_frame = tk.Frame(self.window, bg="#1a1a2e")
        self.main_frame.grid(row=0, column=0, padx=30, pady=30, sticky="nsew")
        self.main_frame.grid_columnconfigure(0, weight=1)
        
        # Create animated background
        self.create_animated_background()
        
        # Title with custom font and icon
        self.title_frame = tk.Frame(self.main_frame, bg="#1a1a2e")
        self.title_frame.grid(row=0, column=0, pady=(40, 20))
        
        # Create PDF icon for title
        self.create_title_icon()
        
        self.title_label = tk.Label(
            self.title_frame,
            text="ASIF's PDF CONVERTER",
            font=("Helvetica", 36, "bold"),
            fg="#e94560",
            bg="#1a1a2e"
        )
        self.title_label.pack(side="right", padx=10)
        
        # Create tab control
        self.tab_control = ttk.Notebook(self.main_frame)
        
        # Create Convert PDF tab with icon
        self.convert_tab = tk.Frame(self.tab_control, bg="#1a1a2e")
        self.tab_control.add(self.convert_tab, text="📄 INSERT PDF FILE")  # Added PDF icon
        
        # Create Recent Files tab with icon
        self.recent_tab = tk.Frame(self.tab_control, bg="#1a1a2e")
        self.tab_control.add(self.recent_tab, text="📋 Recent Files")  # Added list icon
        
        self.tab_control.grid(row=1, column=0, sticky="ew", pady=(0, 20))
        
        # Configure tab style
        style = ttk.Style()
        style.configure("TNotebook", background="#1a1a2e", borderwidth=0)
        style.configure("TNotebook.Tab", 
                       background="#16213e", 
                       foreground="white", 
                       padding=[10, 5],
                       font=("Helvetica", 10, "bold"))
        style.map("TNotebook.Tab",
                 background=[("selected", "#e94560")],
                 foreground=[("selected", "white")])
        
        # Setup Convert PDF tab content
        self.setup_convert_tab()
        
        # Setup Recent Files tab content
        self.setup_recent_tab()
        
        # Initialize recent files list
        self.recent_files = []
        self.max_recent_files = 10
        
        self.selected_file = None
        self.animation_running = True
        self.start_background_animation()

    def create_app_icon(self):
        # Create a 32x32 icon
        icon_size = 32
        icon = Image.new('RGBA', (icon_size, icon_size), (0, 0, 0, 0))
        draw = ImageDraw.Draw(icon)
        
        # Draw PDF icon
        # Background
        draw.rectangle([2, 2, icon_size-2, icon_size-2], fill="#e94560")
        # PDF text
        draw.text((8, 8), "PDF", fill="white", font=None, font_size=12)
        
        # Convert to PhotoImage
        self.app_icon = ImageTk.PhotoImage(icon)
        self.window.iconphoto(True, self.app_icon)

    def create_title_icon(self):
        # Create a 48x48 icon for the title
        icon_size = 48
        icon = Image.new('RGBA', (icon_size, icon_size), (0, 0, 0, 0))
        draw = ImageDraw.Draw(icon)
        
        # Draw PDF icon
        # Background
        draw.rectangle([2, 2, icon_size-2, icon_size-2], fill="#e94560")
        # PDF text
        draw.text((8, 8), "PDF", fill="white", font=None, font_size=20)
        
        # Convert to PhotoImage
        self.title_icon = ImageTk.PhotoImage(icon)
        
        # Create label for icon
        icon_label = tk.Label(
            self.title_frame,
            image=self.title_icon,
            bg="#1a1a2e"
        )
        icon_label.pack(side="left", padx=10)

    def setup_convert_tab(self):
        # Description with icon
        self.desc_label = tk.Label(
            self.convert_tab,
            text="📄 Convert your PDF to Word with perfect layout and font preservation\nSupports English, Arabic, and Urdu languages",
            font=("Helvetica", 14),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.desc_label.pack(pady=(20, 40))
        
        # File selection frame
        self.file_frame = tk.Frame(self.convert_tab, bg="#1a1a2e")
        self.file_frame.pack(fill="x", padx=20, pady=10)
        
        self.file_label = tk.Label(
            self.file_frame,
            text="📄 No file selected",
            font=("Helvetica", 12),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.file_label.pack(side="left", padx=10, pady=10)
        
        self.select_button = tk.Button(
            self.file_frame,
            text="📂 Select PDF",
            command=self.select_file,
            font=("Helvetica", 14, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=2,
            relief=tk.RAISED,
            borderwidth=3
        )
        self.select_button.pack(side="right", padx=10, pady=10)
        
        # Convert button with icon
        self.convert_button = tk.Button(
            self.convert_tab,
            text="🔄 Convert to Word",
            command=self.start_conversion,
            state="disabled",
            font=("Helvetica", 14, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=2,
            relief=tk.RAISED,
            borderwidth=3
        )
        self.convert_button.pack(pady=20)
        
        # Progress bar with custom style
        style = ttk.Style()
        style.configure("Custom.Horizontal.TProgressbar",
                       troughcolor="#16213e",
                       background="#e94560",
                       thickness=25)
        
        self.progress_bar = ttk.Progressbar(
            self.convert_tab,
            orient="horizontal",
            length=400,
            mode="determinate",
            style="Custom.Horizontal.TProgressbar"
        )
        self.progress_bar.pack(pady=10)
        
        # Status label with icon
        self.status_label = tk.Label(
            self.convert_tab,
            text="📄 Ready to convert your PDF",
            font=("Helvetica", 12),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.status_label.pack(pady=10)

    def setup_recent_tab(self):
        # Create a frame for the recent files list
        self.recent_frame = tk.Frame(self.recent_tab, bg="#1a1a2e")
        self.recent_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Create a label for the list with icon
        self.recent_label = tk.Label(
            self.recent_frame,
            text="📋 Recently Converted Files",
            font=("Helvetica", 16, "bold"),
            fg="#ffffff",
            bg="#1a1a2e"
        )
        self.recent_label.pack(pady=(0, 20))
        
        # Create a listbox for recent files
        self.recent_listbox = tk.Listbox(
            self.recent_frame,
            bg="#16213e",
            fg="#ffffff",
            font=("Helvetica", 12),
            selectbackground="#e94560",
            selectforeground="#ffffff",
            height=10
        )
        self.recent_listbox.pack(fill="both", expand=True)
        
        # Add scrollbar
        scrollbar = tk.Scrollbar(self.recent_listbox)
        scrollbar.pack(side="right", fill="y")
        self.recent_listbox.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.recent_listbox.yview)
        
        # Add buttons frame
        button_frame = tk.Frame(self.recent_frame, bg="#1a1a2e")
        button_frame.pack(fill="x", pady=20)
        
        # Add Open File button with icon
        self.open_button = tk.Button(
            button_frame,
            text="📂 Open File",
            command=self.open_selected_file,
            font=("Helvetica", 12, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=1,
            relief=tk.RAISED,
            borderwidth=2
        )
        self.open_button.pack(side="left", padx=5)
        
        # Add Clear List button with icon
        self.clear_button = tk.Button(
            button_frame,
            text="🗑️ Clear List",
            command=self.clear_recent_files,
            font=("Helvetica", 12, "bold"),
            bg="#e94560",
            fg="#ffffff",
            activebackground="#16213e",
            activeforeground="#ffffff",
            width=15,
            height=1,
            relief=tk.RAISED,
            borderwidth=2
        )
        self.clear_button.pack(side="right", padx=5)

    def add_to_recent_files(self, file_path):
        # Add file to recent files list
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        file_info = f"📄 {timestamp} - {os.path.basename(file_path)}"
        
        if file_info not in self.recent_files:
            self.recent_files.insert(0, file_info)
            if len(self.recent_files) > self.max_recent_files:
                self.recent_files.pop()
            
            # Update listbox
            self.recent_listbox.delete(0, tk.END)
            for file in self.recent_files:
                self.recent_listbox.insert(tk.END, file)

    def open_selected_file(self):
        selection = self.recent_listbox.curselection()
        if selection:
            file_info = self.recent_listbox.get(selection[0])
            file_name = file_info.split(" - ")[1]
            file_path = os.path.join(os.path.dirname(self.selected_file), file_name)
            
            if os.path.exists(file_path):
                os.startfile(file_path)
            else:
                messagebox.showerror("Error", "❌ File not found!")

    def clear_recent_files(self):
        self.recent_files.clear()
        self.recent_listbox.delete(0, tk.END)

    def select_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("PDF files", "*.pdf")],
            title="Select PDF File"
        )
        if file_path:
            self.selected_file = file_path
            self.file_label.configure(text=f"📄 {os.path.basename(file_path)}")
            self.convert_button.configure(state="normal")
            self.status_label.configure(text="📄 File selected. Ready to convert!")

    def convert_pdf_to_word(self):
        try:
            self.status_label.configure(text="🔄 Converting your PDF to Word...")
            self.progress_bar["value"] = 20
            
            # Create output filename
            output_file = os.path.splitext(self.selected_file)[0] + ".docx"
            
            # Create a new Word document
            doc = Document()
            
            # Add title to the document
            title = doc.add_heading('Converted PDF Document', 0)
            title.alignment = 1  # Center alignment
            
            # Open the PDF file
            with open(self.selected_file, 'rb') as file:
                # Create PDF reader object
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Get number of pages
                num_pages = len(pdf_reader.pages)
                
                # Process each page
                for page_num in range(num_pages):
                    # Update progress
                    progress = 20 + (80 * (page_num + 1) / num_pages)
                    self.progress_bar["value"] = progress
                    
                    # Get the page
                    page = pdf_reader.pages[page_num]
                    
                    # Extract text
                    text = page.extract_text()
                    
                    # Handle Arabic/Urdu text
                    if any(ord(c) > 127 for c in text):
                        text = get_display(arabic_reshaper.reshape(text))
                    
                    # Add text to document
                    doc.add_paragraph(text)
                    
                    # Add page break if not the last page
                    if page_num < num_pages - 1:
                        doc.add_page_break()
            
            # Save the document
            doc.save(output_file)
            
            # Add to recent files
            self.add_to_recent_files(output_file)
            
            self.progress_bar["value"] = 100
            self.status_label.configure(text="✅ Conversion completed successfully!")
            messagebox.showinfo("Success", f"PDF has been converted to Word successfully!\nSaved as: {os.path.basename(output_file)}")
            
        except Exception as e:
            self.status_label.configure(text="❌ Error during conversion")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
        
        finally:
            self.progress_bar["value"] = 0
            self.convert_button.configure(state="normal")

    def start_conversion(self):
        self.convert_button.configure(state="disabled")
        # Start conversion in a separate thread
        thread = threading.Thread(target=self.convert_pdf_to_word)
        thread.daemon = True
        thread.start()

    def create_animated_background(self):
        # Create a canvas for the animated background
        self.canvas = tk.Canvas(self.window, highlightthickness=0, bg="#1a1a2e")
        self.canvas.place(x=0, y=0, relwidth=1, relheight=1)
        
        # Create gradient circles
        self.circles = []
        for _ in range(5):
            x = self.window.winfo_width() * 0.5
            y = self.window.winfo_height() * 0.5
            circle = self.canvas.create_oval(x-100, y-100, x+100, y+100, 
                                          fill="#e94560", outline="")
            self.circles.append({"id": circle, "dx": 2, "dy": 2})

    def start_background_animation(self):
        def animate():
            if not self.animation_running:
                return
            
            for circle in self.circles:
                # Move circle
                self.canvas.move(circle["id"], circle["dx"], circle["dy"])
                
                # Get current position
                pos = self.canvas.coords(circle["id"])
                
                # Bounce off walls
                if pos[0] <= 0 or pos[2] >= self.window.winfo_width():
                    circle["dx"] *= -1
                if pos[1] <= 0 or pos[3] >= self.window.winfo_height():
                    circle["dy"] *= -1
                
                # Change color gradually
                color = self.canvas.itemcget(circle["id"], "fill")
                if color == "#e94560":
                    self.canvas.itemconfig(circle["id"], fill="#16213e")
                else:
                    self.canvas.itemconfig(circle["id"], fill="#e94560")
            
            self.window.after(50, animate)
        
        animate()

    def run(self):
        self.window.mainloop()
        self.animation_running = False

if __name__ == "__main__":
    app = PDFConverterApp()
    app.run() 